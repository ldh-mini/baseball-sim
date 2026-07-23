"""프로젝트 개요서 → Word(.docx) 변환 스크립트"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '맑은 고딕'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 페이지 여백
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_title(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
    # 데이터
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = '맑은 고딕'
    doc.add_paragraph()  # 테이블 후 간격

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ════════════════════════════════════════════════
# 문서 작성 시작
# ════════════════════════════════════════════════

# 표지
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')
run.font.size = Pt(28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('KBO 야구 경기 결과 예측 서비스')
run.bold = True
run.font.size = Pt(28)
run.font.name = '맑은 고딕'
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('프로젝트 개요서')
run.bold = True
run.font.size = Pt(20)
run.font.name = '맑은 고딕'
run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nv8.1 | 2026-03-27')
run.font.size = Pt(14)
run.font.name = '맑은 고딕'
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('프로토타입 개발 완료 + 예측 고도화 + Statiz 실데이터 반영')
run.font.size = Pt(11)
run.font.name = '맑은 고딕'
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── 1. 프로젝트 비전 ──
add_title('1. 프로젝트 비전')
add_para('**"AI 동반자형 야구 분석 서비스"** — 사용자가 혼자 판단하지 않고, AI가 옆에서 함께 분석해주는 경험을 제공한다.')

add_title('핵심 가치', level=2)
add_bullet('생성형 AI 기반 스포츠 분석 서비스로 **시장 선점** (유사 서비스 미존재)')
add_bullet('"AI가 인간의 분석을 뛰어넘는다" — 알파고 사례와 같은 **강력한 마케팅 스토리**')
add_bullet('초기 무료 분석 제공 → 적중률로 신뢰 확보 → 자연스러운 수익화')

# ── 2. 비즈니스 모델 ──
add_title('2. 비즈니스 모델')

add_title('B2C (소비자 대상)', level=2)
add_table(['단계', '내용'], [
    ['무료 체험', '기본 승리팀 예측 + AI 대화형 분석 무료 제공'],
    ['신뢰 형성', '적중률 검증 데이터 공개로 서비스 신뢰도 확보'],
    ['구독 전환', '프리미엄 분석 (상세 매치업, 백테스트 등) 구독형 판매'],
])

add_title('B2B (기업 대상)', level=2)
add_table(['대상', '제공 가치'], [
    ['KBO 구단', '시뮬레이션 결과 데이터 납품, 전략 지원'],
    ['스포츠 미디어', '경기 예측 콘텐츠 API 제공'],
    ['데이터 분석 업체', '능력치 산정 레시피(자체 모델) 라이센싱'],
])

add_title('확장 계획', level=2)
add_bullet('야구(KBO/MLB) 검증 완료 후 → 축구 등 타 종목으로 확장')
add_bullet('카메라 기능: 경기 화면 촬영 → 즉시 AI 분석 (바이럴 마케팅 효과)')

# ── 3. 기술 아키텍처 ──
add_title('3. 기술 아키텍처')

add_title('기술 스택', level=2)
add_table(['구분', '기술'], [
    ['프론트엔드', 'React 19 + Vite 8 + Tailwind CSS 3'],
    ['UI 테마', '다크 블루/퍼플 글래스모피즘 (deepbetting.io 참고)'],
    ['시뮬레이션 엔진', 'JavaScript (클라이언트 사이드)'],
    ['통계 방법론', '몬테카를로 시뮬레이션 (100~10,000회) + 세이버메트릭스 고급 지표'],
    ['데이터 소스', 'Statiz 2025 실데이터 (WAR/FIP 포함) + KBO 공식 일정 API'],
    ['경기 일정', 'KBO 공식 API 실시간 연동 (Vite 프록시 CORS 우회)'],
    ['날씨 API', 'wttr.in (실시간 날씨 연동)'],
    ['AI 분석', '키워드 파싱 기반 인텔리전스 시스템'],
])

add_title('파일 구조', level=2)
add_code_block("""야구시뮬레이션/
  kbo-simulation.jsx   # 전체 앱 (엔진 + UI, ~900줄)
  src/main.jsx         # 엔트리포인트
  src/index.css        # 글로벌 다크 테마 스타일
  index.html           # HTML 템플릿
  package.json         # 의존성 관리
  vite.config.js       # 빌드 설정 + KBO API 프록시
  tailwind.config.js   # 커스텀 다크 테마
  backtest-runner.mjs  # 2025 시즌 백테스트 (독립 실행)
  backtest-2024.mjs    # 2024 시즌 백테스트 (독립 실행)""")

# ── 4. 시뮬레이션 엔진 상세 ──
add_title('4. 시뮬레이션 엔진 상세')

add_title('4.1 데이터 수록 범위', level=2)
add_table(['구분', '내용', '수록량'], [
    ['팀', 'KBO 10개 구단 전체', '10팀'],
    ['타자', '팀당 9명, Statiz 2025 실데이터', '90명 (WAR 보유 60+명)'],
    ['투수', '팀당 선발 3명 + 불펜, Statiz 실데이터', '30명 (전원 WAR/FIP) + 불펜'],
    ['구장', '9개 구장 (파크팩터/돔 여부)', '9곳'],
    ['날씨', '6종 (맑음/흐림/비/추위/더위/강풍)', '6종'],
    ['백테스트', '2025 시즌 실제 경기 결과', '60경기'],
])

add_title('4.2 반영 변수 (28개 카테고리)', level=2)

add_title('선수 개인 스탯', level=3)
add_table(['변수', '타자', '투수'], [
    ['기본 스탯', 'AVG, OBP, SLG, HR, SPD', 'ERA, WHIP, K/9, BB/9, IP'],
    ['고급 스탯', 'WAR, defRAA, totalAvg, RBI, wOBA', 'WAR, WPA/LI, FIP, 투구엔트로피'],
    ['컨디션', 'recentForm (최근 폼)', 'recentForm (최근 폼)'],
    ['좌우 매치업', 'bat (L/R/S)', 'throws (L/R)'],
])

add_title('환경 요소', level=3)
add_table(['변수', '설명', '보정 범위'], [
    ['구장 파크팩터', '사직 1.08 ~ 고척 0.92', '안타/홈런 확률'],
    ['돔구장 여부', '날씨 영향 80% 감소', '날씨 보정 축소'],
    ['날씨', '안타/홈런/실책 확률 보정', 'hitMod, hrMod, errMod'],
    ['홈어드밴티지', '홈팀 +4% 전체 보정', '타석 확률'],
])

add_title('v6.0 보정 (5개)', level=3)
add_table(['변수', '설명', '보정 범위'], [
    ['요일별 성적', '주말 홈팀 유리, 월요일 부진', '홈 0.97~1.04 / 원정 0.97~1.01'],
    ['시간대 성적', '주간 타고투저, 야간 투고타저', '안타 0.98~1.04, HR 0.97~1.06'],
    ['배당값 보정', '언더독 부스트 / 탑독 페널티', '언더독 최대+5%, 탑독 최대-3%'],
    ['팀 상대전적(H2H)', '2025 시즌 팀간 승률', '승률 .600 → +1.5%'],
    ['투수-타자 매치업', '통산 상대 타율 기반', '최대 +/-8%'],
])

add_title('v7.0 신규 기능 (3개)', level=3)
add_table(['변수', '설명', '보정 범위'], [
    ['투수 피로도', '이닝/실점/피안타 누적 → 능력 하락', '삼진-50%, 볼넷+60%, 피안타+80%'],
    ['지능형 투수 교체', '피로도/실점/스코어/에이스 종합 판단', '7회 일괄 → 상황별 교체'],
    ['도루/희생번트', '속력 기반 도루, 접전 시 하위타선 번트', 'SPD 7+ 도루, 7회+ 번트'],
])

add_title('v8.0 신규 기능 (10개)', level=3)
add_table(['기능', '설명'], [
    ['wOBA (가중출루율)', '타자 통합 공격 지표 — AVG/OBP/SLG 개별 대신 단일 지표로 정확도 향상'],
    ['FIP (수비무관방어율)', 'ERA 대신 투수 본연의 능력만 측정 — 홈런/볼넷/삼진 기반, 수비 노이즈 제거'],
    ['피타고리안 기대승률', 'RS^1.83 / (RS^1.83 + RA^1.83) — 팀 득실점 기반 실제 전력 (KBO 최적화 지수 1.83)'],
    ['Elo 동적 레이팅', '실제 승률(60%) + 피타고리안(40%) 가중 블렌딩 → 팀 전력 동적 보정 (±6%)'],
    ['평균 회귀 보정', 'PA/IP 표본 크기 기반 소표본 보정 — 리그 평균으로 자동 회귀 (과적합 방지)'],
    ['리그 평균 상수', 'KBO 2025 리그 평균값 (AVG .265, ERA 3.80 등) 기준점으로 활용'],
    ['KBO 공식 일정 연동', 'koreabaseball.com API → 실제 날짜별 경기 일정 실시간 반영'],
    ['과거 경기 결과 표시', '과거 날짜 선택 시 실제 스코어 + 승패 뱃지 표시'],
    ['오늘의 경기 인라인 시뮬', '경기 카드 클릭 → 확장 레이어에서 바로 채팅 + 시뮬레이션'],
    ['다크 블루/퍼플 UI', '글래스모피즘 + 네온 글로우 + 그라디언트 기반 전문 서비스 UI'],
    ['Statiz 2025 실데이터', '전 구단 선수 스탯을 Statiz/FancyStats 기반 실측치로 교체 (WAR/FIP 전원 반영)'],
    ['로스터 최신화', 'MLB 이적(이정후/김혜성), 트레이드(손아섭), 부상(안우진), 은퇴(추신수) 등 7건 반영'],
    ['KBO 공식 구단 로고', 'NaverCDN 실제 엠블럼 이미지 적용 (이모지 → 공식 로고)'],
    ['외부 접속 지원', 'Cloudflared 터널링으로 모바일 외부 테스트 환경 구축'],
])

add_title('4.3 확률 계산 흐름', level=2)
add_code_block("""[v8.0 사전 처리]
  타자 → 평균 회귀(PA 기반) → wOBA 계산
  투수 → 평균 회귀(IP 기반) → FIP 계산
  팀   → 피타고리안 기대승률 → Elo 레이팅 → eloMod 보정

[매 타석 확률 계산]
  타자력 = wOBA / 리그평균(.340) x 홈어드밴티지 x 날씨 x 폼 x 좌우 x WAR x 환경보정 x 피로보정
  투수력 = (4.5 - FIP) / 4.5 + 0.5 x 폼 x WPA x 시간대보정

  환경보정(envMod) = 요일보정 x 배당보정 x H2H보정 x 매치업보정 x Elo보정

  결과 분류: 안타 / 2루타 / 3루타 / 홈런 / 볼넷 / 사구 / 삼진 / 땅볼 / 뜬공 / 라인아웃 / 실책""")

add_title('4.4 게임 진행 로직', level=2)
add_bullet('정규 9이닝 + 연장 최대 12이닝')
add_bullet('[v7.0] 지능형 투수 교체 (피로도/실점/스코어 종합 판단)')
add_bullet('[v7.0] 이닝별 투수 피로도 반영 (삼진 하락, 볼넷/피안타 증가)')
add_bullet('[v7.0] 도루 시도 (SPD 7+ 주자, 접전 시 확률 증가)')
add_bullet('[v7.0] 희생번트 (7회+ 하위타선, 접전 시 주자 진루)')
add_bullet('끝내기 처리 (9회 이후 홈팀 리드 시)')
add_bullet('주루 판단 (속력 기반 진루/득점 확률)')
add_bullet('병살, 희생플라이 확률 적용')
add_bullet('이닝 중 3실점 이상 시 긴급 투수 교체')

# ── 5. UI/UX ──
add_title('5. UI/UX 구성')

add_title('5.1 오늘의 경기 (TodayTab)', level=2)
add_bullet('날짜 선택 → **KBO 공식 API에서 실제 경기 일정 로드** (Vite 프록시 CORS 우회)')
add_bullet('월별 캐시로 같은 달 재요청 방지')
add_bullet('"KBO 공식 일정" / "자동 생성 일정" 데이터 소스 뱃지 표시')
add_bullet('**과거 경기**: 실제 스코어 표시 (이긴 팀 시안색 강조 + 승패 뱃지)')
add_bullet('**예정 경기**: "VS" 표시')
add_bullet('**인라인 시뮬레이션**: 경기 카드 클릭 → 아래로 확장')
add_bullet('AI 채팅 (BALL-E): 선수 컨디션 자연어 입력 → 수치 변환', level=1)
add_bullet('구장 정보 + 상대전적(H2H) 패널', level=1)
add_bullet('단일 경기 / 몬테카를로 시뮬레이션 (100~10,000회)', level=1)
add_bullet('승률 바, 득점 분포, 스코어보드, 이닝별 로그', level=1)
add_bullet('실시간 날씨 API 연동 (wttr.in)')
add_bullet('경기 없는 날: "해당 날짜에 예정된 경기가 없습니다" 안내')

add_title('5.2 가상 대결 (VirtualTab)', level=2)
add_bullet('홈/원정팀 자유 선택 (10팀), 선발투수 선택 (팀당 3명)')
add_bullet('요일/시간대/날씨 수동 설정')
add_bullet('**AI 인텔리전스 채팅**: 자연어로 선수 컨디션 입력 → 수치 자동 변환')
add_bullet('단일 경기 시뮬레이션 + 몬테카를로 (100~10,000회)')
add_bullet('승률 비교, 득점 분포, 이닝별 상세 로그')

add_title('5.3 백테스트 (BacktestTab)', level=2)
add_bullet('2025 시즌 실제 경기 60건 vs 시뮬레이션 예측 비교')
add_bullet('전체 적중률 산출')
add_bullet('경기별 예측 승률, 예측 스코어, 실제 결과, 적중 여부 테이블')

add_title('5.4 UI/디자인 (v8.0)', level=2)
add_bullet('다크 블루/퍼플 테마 (deepbetting.io 참고)')
add_bullet('글래스모피즘 카드 (backdrop-filter blur + 반투명 그라디언트)')
add_bullet('네온 글로우 효과 (텍스트/보더 퍼플/블루 글로우)')
add_bullet('그라디언트 버튼 (블루→퍼플 + hover 애니메이션)')
add_bullet('반응형 레이아웃 (모바일/데스크톱)')

# ── 6. 마케팅 전략 ──
add_title('6. 마케팅 전략')
add_table(['전략', '내용'], [
    ['핵심 메시지', '"AI가 인간의 분석을 뛰어넘는다"'],
    ['신뢰 확보', '초기 무료 분석 → 적중률 공개 → 자연스러운 결제 유도'],
    ['차별화', 'AI 캐릭터(BALL-E)와 대화하며 분석하는 동반자형 경험'],
    ['바이럴', '카메라 촬영 → 즉시 분석 기능 (향후 구현)'],
    ['벤치마킹', 'deepbetting.io (디자인 참고), 메이저리그 스몰볼 성공사례'],
])

# ── 7. 현재 상태 vs 목표 ──
add_title('7. 현재 구현 상태 vs 최종 목표')
add_table(['기능', '현재 상태', '최종 목표'], [
    ['시뮬레이션 엔진', '28개 변수 + 세이버메트릭스 + 피로도/교체AI', '서버사이드 + 실시간 데이터'],
    ['데이터', 'Statiz 2025 실데이터 (WAR/FIP 전원) + KBO 공식 일정 API', '실시간 크롤링/API 자동 갱신'],
    ['경기 일정', 'KBO 공식 API 실시간 연동 (완료)', '선수 데이터도 자동 갱신'],
    ['과거 결과', '실제 스코어 표시 (완료)', '경기별 상세 기록 연동'],
    ['AI 분석', '키워드 파싱 (InsightParser)', '생성형 AI (LLM) 연동'],
    ['투수 교체', 'v7.0 상황별 지능형 교체 완료', '중계/마무리 분리'],
    ['UI/디자인', 'v8.0 다크 블루/퍼플 글래스모피즘', '인터랙티브 차트/3D 시각화'],
    ['백테스트', 'v8.1: 75.0% (고신뢰 78.4%) Statiz 실데이터', '전체 시즌 (720경기+)'],
    ['사용자 인증', '없음', '로그인/구독 시스템'],
    ['결제', '없음', '구독형 결제 연동'],
    ['종목', 'KBO 야구만', 'MLB + 축구 등 확장'],
    ['카메라 분석', '미구현', '화면 촬영 → 즉시 분석'],
    ['모바일 앱', '웹 전용', '네이티브 앱 (iOS/Android)'],
])

# ── 8. 로드맵 ──
add_title('8. 향후 개발 로드맵')

add_title('Phase 1 — 엔진 고도화 (완료)', level=2)
for item in ['요일별/시간대별 성적 보정', '배당값(언더독/탑독) 보정', '선발투수-타자 상대전적 매치업',
             '팀 상대전적(H2H) 반영', '백테스트 탭 (2025 시즌 60경기)',
             '투수 교체 전략 지능화 (v7.0)', '도루/희생번트 전략 (v7.0)', '이닝별 투수 피로도 반영 (v7.0)']:
    add_bullet(f'[완료] {item}')

add_title('Phase 1.5 — UI/UX + 데이터 연동 (완료)', level=2)
for item in ['다크 블루/퍼플 UI 리디자인 (글래스모피즘/네온 글로우)',
             'KBO 공식 일정 API 연동 (실제 날짜별 경기 일정)',
             '과거 경기 실제 결과(스코어) 표시',
             '오늘의 경기 인라인 시뮬레이션 (경기 클릭 → 확장 레이어)',
             '오늘의 경기 탭에 AI 채팅 + 시뮬레이션 옵션 통합',
             '2024/2025 시즌 독립 백테스트 스크립트 작성',
             '백테스트 검증: 2024 73.3%, 2025 71.7% (고신뢰 ~78%)',
             'Vite 프록시 설정 (KBO API CORS 우회)']:
    add_bullet(f'[완료] {item}')

add_title('Phase 1.7 — 예측 고도화 (완료)', level=2)
for item in ['wOBA (가중출루율) 도입 — 타자 통합 공격 지표로 예측력 강화',
             'FIP (수비무관방어율) 도입 — ERA 대체, 투수 본연 능력 측정',
             '피타고리안 기대승률 — 팀 득실점 기반 전력 지표 (지수 1.83 KBO 최적화)',
             'Elo 동적 레이팅 — 승률 + 피타고리안 가중 블렌딩, 팀 전력 동적 보정',
             '평균 회귀 보정 — PA/IP 표본 크기 기반 소표본 과적합 방지',
             '리그 평균 상수 — KBO 2025 기준값 설정 (AVG .265, ERA 3.80 등)',
             '백테스트 재검증: v7.0 71.7% → v8.0 76.7% (+5.0%p 향상)',
             '고신뢰(60%+) 적중률: 80.4% (51경기 중 41경기 적중)']:
    add_bullet(f'[완료] {item}')

add_title('Phase 1.8 — Statiz 실데이터 + 비주얼 고도화 (완료)', level=2)
for item in ['Statiz 2025 시즌 실데이터 반영 — 전 구단 90명 타자 + 30명 투수 실측 스탯',
             '투수 WAR/FIP 전원 반영 (기존 5/30 → 30/30)',
             '타자 WAR 60+명 반영 (기존 ~20% → ~70%)',
             '로스터 최신화 7건: MLB 이적, 트레이드, 부상, 은퇴, 외국인 교체',
             '투구손 교정: 임찬규/김광현/김윤식 → 좌투(L)',
             'KBO 공식 구단 로고 적용 (NaverCDN 엠블럼)',
             'Cloudflared 터널링 외부 접속 환경 구축',
             '백테스트 결과: v8.1 75.0% (45/60), 고신뢰 78.4% (40/51)']:
    add_bullet(f'[완료] {item}')

add_title('Phase 2 — 데이터 자동화', level=2)
add_bullet('[완료] KBO 공식 경기 일정 API 자동화')
add_bullet('[예정] Statiz 선수 스탯 크롤링 자동화')
add_bullet('[예정] 실시간 선수 컨디션 반영')
add_bullet('[예정] 백테스트 전체 시즌 확장 (720경기+)')
add_bullet('[예정] 서버사이드 시뮬레이션 이전')

add_title('Phase 3 — AI 고도화', level=2)
add_bullet('[예정] 생성형 AI(LLM) 연동 → 자연어 대화형 분석')
add_bullet('[예정] AI 캐릭터(BALL-E) 페르소나 강화')
add_bullet('[예정] 카메라 촬영 → 경기 화면 즉시 분석')

add_title('Phase 4 — 수익화', level=2)
add_bullet('[예정] 사용자 인증 + 구독 시스템')
add_bullet('[예정] B2C 프리미엄 분석 상품 설계')
add_bullet('[예정] B2B 구단/미디어 대상 데이터 납품 체계')

add_title('Phase 5 — 확장', level=2)
add_bullet('[예정] MLB 데이터 추가')
add_bullet('[예정] 타 종목(축구 등) 확장')
add_bullet('[예정] 모바일 네이티브 앱')

# ── 9. 백테스트 검증 결과 ──
doc.add_page_break()
add_title('9. 백테스트 검증 결과')

add_title('9.1 버전별 적중률 비교', level=2)
add_table(['버전', '전체 적중률', '고신뢰(60%+)', '주요 개선'], [
    ['v7.0', '71.7% (43/60)', '~78%', '투수 피로도, 지능형 교체, 도루/번트'],
    ['v8.0', '76.7% (46/60)', '80.4% (41/51)', 'wOBA, FIP, 피타고리안, Elo, 평균 회귀'],
    ['v8.1', '75.0% (45/60)', '78.4% (40/51)', 'Statiz 실데이터, 로스터 최신화, WAR/FIP 전원'],
])
add_para('**v8.0 → v8.1 변동 (-1.7%p)**: 임의 추정치를 실측 데이터로 교체하면서 일부 과적합이 제거된 결과. '
         '실데이터 기반이므로 실전 예측력은 v8.1이 더 신뢰할 수 있음.')

add_title('9.2 v8.1 월별 적중률', level=2)
add_table(['월', '적중률', '적중/전체'], [
    ['3월', '60.0%', '6/10'],
    ['4월', '70.0%', '7/10'],
    ['5월', '70.0%', '7/10'],
    ['6월', '80.0%', '8/10'],
    ['7월', '80.0%', '4/5'],
    ['8월', '100.0%', '5/5'],
    ['9월', '80.0%', '8/10'],
])

add_title('9.3 v8.1 팀별 적중률', level=2)
add_table(['팀', '적중률', '적중/전체'], [
    ['LG', '91.7%', '11/12'],
    ['기아', '83.3%', '10/12'],
    ['한화', '83.3%', '10/12'],
    ['SSG', '83.3%', '10/12'],
    ['두산', '75.0%', '9/12'],
    ['KT', '75.0%', '9/12'],
    ['NC', '75.0%', '9/12'],
    ['키움', '75.0%', '9/12'],
    ['삼성', '58.3%', '7/12'],
    ['롯데', '50.0%', '6/12'],
])

add_title('9.4 데이터 품질 개선 (v8.1)', level=2)
add_table(['항목', 'v8.0 이전', 'v8.1 이후'], [
    ['투수 WAR 보유', '5/30명 (17%)', '30/30명 (100%)'],
    ['투수 FIP 보유', '3/30명 (10%)', '25/30명 (83%)'],
    ['타자 WAR 보유', '~18/90명 (20%)', '60+/90명 (67%)'],
    ['로스터 정확도', 'MLB 이적자 포함', '최신 로스터 반영'],
    ['투구손 정확도', '좌투 3명 오류', '전원 교정 완료'],
    ['팀 실적 데이터', '추정치', '2025 시즌 실제 성적'],
])

add_title('9.5 예측 이론적 한계', level=2)
add_para('야구 경기 예측의 이론적 상한선은 약 **58~62%** (MLB 기준, 대규모 베팅 시장 데이터). '
         '현재 75.0%는 60경기 소규모 샘플 + 시즌 종합 스탯 기반(look-ahead 포함)이므로, '
         '실전 배포 시 **실시간 데이터 + 대규모 샘플(720경기+)** 으로 재검증 필요.')

# ── 10. 참고 자료 ──
add_title('10. 참고 자료')
add_bullet('디자인 참고: deepbetting.io')
add_bullet('사업 근거 영상:')
add_bullet('youtu.be/SQlHaWT9YTE', level=1)
add_bullet('youtu.be/eRVry274WAc', level=1)
add_bullet('youtu.be/37Eg8KlYCtY', level=1)
add_bullet('youtube.com/watch?v=KkjD68AY7dc', level=1)

# 푸터
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('본 문서는 2026-03-21 1차 미팅 내용 및 프로토타입 코드 분석을 기반으로 작성되었습니다.')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '맑은 고딕'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('v8.1 업데이트: 2026-03-27 — Statiz 2025 실데이터 전면 반영, 로스터 최신화 7건, KBO 공식 로고, 백테스트 75.0%')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '맑은 고딕'

# ── 저장 ──
output_path = r'c:\Users\uniun\Documents\Claude\Projects\야구시뮬레이션\KBO_야구_예측_서비스_프로젝트_개요서_v8.1.docx'
doc.save(output_path)
print(f'Word 파일 생성 완료: {output_path}')
